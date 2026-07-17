import io
import docx
import fitz
import cv2
import numpy as np
import pytesseract
import os
from pypdf import PdfWriter, PdfReader
from PIL import Image
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from deep_translator import GoogleTranslator
from langdetect import detect
from gtts import gTTS

# --- SYSTEM ENVIRONMENT PATH FIX FOR CLOUD DEPLOYMENT ---
if os.path.exists("/usr/bin/tesseract"):
    pytesseract.pytesseract.tesseract_cmd = "/usr/bin/tesseract"

def convert_docx_to_txt(file_bytes):
    doc = docx.Document(io.BytesIO(file_bytes))
    return "\n".join([p.text for p in doc.paragraphs])

def text_to_pdf(text):
    pdf_buffer = io.BytesIO()
    c = canvas.Canvas(pdf_buffer, pagesize=letter)
    width, height = letter
    textobject = c.beginText(50, height - 50)
    textobject.setFont("Helvetica", 11)
    for line in text.split("\n"):
        textobject.textLine(line)
        if textobject.getY() < 50:
            c.drawText(textobject)
            c.showPage()
            textobject = c.beginText(50, height - 50)
            textobject.setFont("Helvetica", 11)
    c.drawText(textobject)
    c.save()
    return pdf_buffer.getvalue()

def text_to_docx(text):
    doc = docx.Document()
    for line in text.split("\n"):
        if line.strip():
            doc.add_paragraph(line)
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    return doc_buffer.getvalue()

def merge_documents(files_list, add_cover=False):
    writer = PdfWriter()
    if add_cover:
        cover_io = io.BytesIO()
        c = canvas.Canvas(cover_io, pagesize=letter)
        c.setFont("Helvetica-Bold", 24)
        c.drawCentredString(300, 500, "AI Smart Document Studio")
        c.setFont("Helvetica", 14)
        c.drawCentredString(300, 450, "Merged Document Ledger")
        c.save()
        writer.append(io.BytesIO(cover_io.getvalue()))

    for file_obj in files_list:
        ext = file_obj.name.split('.')[-1].lower()
        b_data = file_obj.getvalue()
        if ext == 'pdf':
            writer.append(io.BytesIO(b_data))
        elif ext in ['jpg', 'jpeg', 'png', 'bmp', 'tiff']:
            img = Image.open(io.BytesIO(b_data))
            img_pdf = io.BytesIO()
            c = canvas.Canvas(img_pdf, pagesize=(img.width, img.height))
            c.drawInlineImage(img, 0, 0, width=img.width, height=img.height)
            c.save()
            writer.append(io.BytesIO(img_pdf.getvalue()))
        elif ext == 'docx':
            writer.append(io.BytesIO(text_to_pdf(convert_docx_to_txt(b_data))))
    out = io.BytesIO()
    writer.write(out)
    return out.getvalue()

def extract_pdf_pages(pdf_bytes, page_numbers):
    reader = PdfReader(io.BytesIO(pdf_bytes))
    writer = PdfWriter()
    for p in page_numbers:
        writer.add_page(reader.pages[p - 1])
    out = io.BytesIO()
    writer.write(out)
    return out.getvalue()

def run_ocr(file_obj):
    ext = file_obj.name.split('.')[-1].lower()
    text = ""
    if ext == 'pdf':
        doc = fitz.open(stream=file_obj.read(), filetype="pdf")
        text = "".join([page.get_text() for page in doc])
    else:
        np_arr = np.frombuffer(file_obj.read(), np.uint8)
        img = cv2.imdecode(np_arr, cv2.IMREAD_COLOR)
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        text = pytesseract.image_to_string(gray)
    return text

def translate_text(text, target_lang):
    try:
        return GoogleTranslator(source='auto', target=target_lang).translate(text)
    except Exception:
        return "Translation failed. Check network connection."

def text_to_speech(text, lang_code):
    tts = gTTS(text=text, lang=lang_code, slow=False)
    audio_io = io.BytesIO()
    tts.write_to_fp(audio_io)
    return audio_io.getvalue()

